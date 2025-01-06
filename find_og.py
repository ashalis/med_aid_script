import os
import openai
import PyPDF2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

openai.api_key = 'yourkeyhere' #insert your key in the quotes

def iterate(path_name):
    """
    Iterates through the files/folders in a directory. For each item, it checks the item type. If the item is a .docx, calls method to find
    year, school, residency type, and location within the .docx file. If it is a .pdf, calls a method that does the same thing except for pdfs.
    If it is a folder, recursively iterates through the contents, thus reaching every file in the directory. If a file is not a .docx or a 
    .pdf, prints "incompatible file type."

    Parameters: path_name- string
    Preconditions: provide a valid path_name. All files within must end with a year and contain school, residency type, and location. 
    
    """
    content=''
    for file in os.listdir(path_name):
        path = path_name+'/'+file
        if os.path.isfile(path):
            #code to search for info
            if path.endswith('.pdf'):
                content+=(find_in_pdf(path))
            elif path.endswith('.docx'):
                content+=(find_in_docx(path))
            else: 
                print("incompatible file type")
            content+=('~~~')
        else:
            content+=iterate(path)
    return(content)
    

def find_in_docx(path):
    """
    Uses an openai engine to get information from a Docx file.

    Parameters: path is a string
    Precondition: Must be a valid pathname

    """
    #gets school and year from file name
    dot = path.find('.')
    year = path[dot-4:dot]
    uscore1 = path.find('_')
    uscore2 = path[uscore1+1:].find('_')+uscore1
    school = path[uscore1+1:uscore2+1]

    document = Document(path)  

    # Extract the text from the Word document
    text = ""
    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"
    
    #use ai to make a query
    query = "Create a list of residency types (as in specialties) and schools in the format: \" Specialty/Residency Type: School, School, School; Specialty/Residency Type: School, School, School\" and so on. Include repeats of schools as seperate entries (not schoolx2 or similar)."
    prompt = f"Search for: {query}\nFile Contents: {text}\nAnswer:"
    #saves response-- uses chat completion
    response = openai.ChatCompletion.create(
        engine='gpt-3.5-turbo',
        prompt=prompt,
        max_tokens=100,
        n=1,
        stop=None,
        temperature=.7
    )
    print("Docx: "+path)
    print(response.choices[0].text)
    result = year+'@@'+school+'##'+response.choices[0].text
    return result
    
    

def find_in_pdf(path):
    """
    Uses an openai engine to get information from a pdf. 

    Parameters: path is a string
    Precondition: Must be a valid pathname

    """
    #gets school and year from file name
    dot = path.find('.')
    year = path[dot-4:dot]
    uscore1 = path.find('_')
    uscore2 = path[uscore1+1:].find('_')+uscore1
    school = path[uscore1+1:uscore2+1]
    with open(path, 'rb') as file:
        reader = PyPDF2.PdfFileReader(file)
        text = ""
        for page in range(reader.numPages):
            text += reader.getPage(page).extractText()

    #use ai to make a query
    query = "Create a list of residency types (as in specialties) and schools in the format: \" Specialty/Residency Type: School, School, School; Specialty/Residency Type: School, School, School\" and so on. Each individual should be a separate entry."
    prompt = f"Search for: {query}\nFile Contents: {text}\nAnswer:"
    #saves response-- uses chat completion
    response = openai.Completion.create(
        engine='text-davinci-003',
        prompt=prompt,
        max_tokens=100,
        n=1,
        stop=None,
        temperature=.7
    )
    print("pdf: "+path)
    result = year+'@@'+school+'##'+response.choices[0].text
    print(result)
    return(result)
    
#I think the print statements in this funciton are all just checks so you can delete them 
def format(content2):
    """
    Function to create and format a table with all the rows and collumns filled with the date from the iterate function.
    Parameters: content2 is a string containing markers for data
    Precondition: content2 must be in the correct format (produced by iterate function)
    """
    ## creates a document with a table with four columns
    doc = Document()
    doc.add_table(1,4)

    row = 1

    #labels columns
    doc.tables[0].cell(0,0).text="Year"
    doc.tables[0].cell(0,1).text="School"
    doc.tables[0].cell(0,2).text="Residency Type"
    doc.tables[0].cell(0,3).text="Location"
    
    #fills columns
    for count in range(content2.count('@@')):
        #Finds year and school marks and stores to fill the table in the inner for loop
        year_mark = content2.find('@@')
        school_mark = content2.find('##')
        year = content2[:year_mark]
        print('year: '+year)
        school = content2[year_mark+2:school_mark]
        print('school: '+school)
        content2 = content2[school_mark+3:]
        residencies = content2[:content2.find('~~~')]
        for count in range(residencies.count(':')):
            #finds where the residencies schools start
            endpt = residencies.find(':')
            residency = residencies[:endpt]
            print(residency)
            #finds end of the residencies
            schools = residencies[endpt+1:residencies.find(';')]
            print(residency+": "+schools)
            print(residencies[endpt+1:])
            for locations in range(schools.count(',')):
                location1=schools[:schools.find(',')]
                schools=schools[schools.find(',')+1:]
                doc.tables[0].add_row()
                
                doc.tables[0].cell(row,0).text=year
                doc.tables[0].cell(row,1).text=school
                doc.tables[0].cell(row,2).text=residency
                doc.tables[0].cell(row,3).text=location1

                row+=1
            doc.tables[0].add_row()
            doc.tables[0].cell(row,0).text=year
            doc.tables[0].cell(row,1).text=school
            doc.tables[0].cell(row,2).text=residency
            doc.tables[0].cell(row,3).text=schools[:]
            row+=1
            residencies = residencies[residencies.find(';')+2:]
        content2=content2[content2.find('~~~')+3:]
            #table.
    # Customize table formatting
    table=doc.tables[0]
    table.autofit = False

    # Set column widths
    column_widths = [Inches(1.5)] * 4
    for i, width in enumerate(column_widths):
        table.columns[i].width = width

    # Set alignment of cells
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add borders to the table
    table.style = "Table Grid"
    doc.save("output2.docx")